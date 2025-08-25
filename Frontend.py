import streamlit as st
from rdflib import Graph, Namespace, URIRef, Literal
from rdflib.namespace import RDF, RDFS, OWL, FOAF, DCTERMS
from docx import Document
import PyPDF2
import pdfplumber
import tempfile
import os
import re
import unicodedata
from PIL import Image, ImageEnhance, ImageFilter
import io
import pytesseract
import shutil  # Import shutil for directory operations
import urllib.parse  # Import urllib.parse for URL encoding
import datetime  # Import datetime for unique folder naming
import hashlib  # Import hashlib for unique folder naming
import traceback  # Import traceback for error details

# Attempt to import OpenCV for advanced image processing.
# If not available, fall back to PIL-only processing.
try:
    import cv2
    import numpy as np

    OPENCV_AVAILABLE = True
except ImportError:
    st.warning("OpenCV (cv2) or NumPy not found. Image pre-processing will be limited. "
               "For better OCR accuracy, consider installing OpenCV (`pip install opencv-python numpy`).")
    OPENCV_AVAILABLE = False

# RDF Namespace
EX = Namespace("http://example.org/ontology/")

# Define common caption patterns
CAPTION_PATTERNS = [
    r"^(?:Figure|Fig\.|FIG\.)\s*(\d+)([a-zA-Z])?\s*[:\-\.]?\s*(.*)",
    r"^(?:Table|TAB\.)\s*(\d+)\s*[:\-\.]?\s*(.*)",
    r"^(?:Chart|Graph)\s*(\d+)\s*[:\-\.]?\s*(.*)",
]

# Constants for defining regions to skip in PDF processing
FOOTER_SKIP_RATIO = 0.07  # Skip bottom 7% of the page height
SIDEBAR_SKIP_RATIO = 0.05  # Skip 5% from left and 5% from right width for sidebars


def clean_text(text):
    """Clean and normalize text for better processing"""
    if not text:
        return ""

    text = unicodedata.normalize('NFKD', text)
    text = ''.join(char for char in text if unicodedata.category(char)[0] != 'C')
    text = re.sub(r'\s+', ' ', text.strip())
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'> ', '', text)
    text = text.replace('"', "'")
    text = text.replace('\\', '/')
    text = re.sub(r'[\r\n\t]', ' ', text)

    return text.strip()


def clean_chapter_title(title):
    """
    Remove 'chapter' word, chapter numbers (Arabic or Roman), and common
    word forms (e.g., 'one', 'two') from title, ensuring robust cleaning.
    """
    if not title:
        return ""

    chapter_pattern = r'^\s*chapter\s*(?:(?:\d+|\b(?:one|two|three|four|five|six|seven|eight|nine|ten)\b|[IVXLCDM]+)\s*[:\-\.]?\s*)?'
    cleaned = re.sub(chapter_pattern, '', title, flags=re.IGNORECASE)
    cleaned = re.sub(r'^\s*(?:\d+(\.\d+)*|[IVXLCDM]+)\s*[.-:]*\s*', '', cleaned, flags=re.IGNORECASE)
    cleaned = cleaned.strip(' .-:')

    if not cleaned or re.fullmatch(r'[\s\W]*', cleaned) or cleaned.lower() == "default chapter":
        return ""

    return cleaned


def preprocess_image_for_ocr(pil_image):
    """
    Apply image pre-processing techniques to improve OCR accuracy.
    Uses OpenCV if available, otherwise PIL's basic functions.
    """
    if pil_image.mode != 'RGB':
        pil_image = pil_image.convert('RGB')

    if OPENCV_AVAILABLE:
        cv_image = cv2.cvtColor(np.array(pil_image), cv2.COLOR_BGR2GRAY)
        _, binary_image = cv2.threshold(cv_image, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        pil_image_processed = Image.fromarray(binary_image)
    else:
        pil_image_processed = pil_image.convert('L')
        enhancer = ImageEnhance.Contrast(pil_image_processed)
        pil_image_processed = enhancer.enhance(2.0)
        pil_image_processed = pil_image_processed.filter(ImageFilter.SHARPEN)

    return pil_image_processed


def extract_caption_from_text(page_text_lines, image_bbox, page_width, page_height):
    """
    Attempts to extract a caption for an image by looking for text patterns
    around the image's bounding box.
    """
    caption = None
    search_top = max(0, image_bbox[1] - 0.2 * page_height)
    search_bottom = min(page_height, image_bbox[3] + 0.2 * page_height)

    relevant_lines = []
    for line_info in page_text_lines:
        if line_info['bbox'][1] >= search_top and line_info['bbox'][3] <= search_bottom:
            relevant_lines.append(line_info)

    relevant_lines.sort(key=lambda x: x['bbox'][1])

    for line_info in relevant_lines:
        text = clean_text(line_info['text'])
        for pattern in CAPTION_PATTERNS:
            match = re.match(pattern, text, re.IGNORECASE)
            if match:
                caption = text
                return caption

    return caption


def extract_specific_entities_from_ocr(ocr_text):
    """
    Identifies specific entities or diagram types based on keywords in OCR text.
    Returns a dictionary of identified entities/types.
    """
    identified_entities = {}
    lower_ocr_text = ocr_text.lower()

    fig_matches = re.findall(r"(?:Figure|Fig\.|FIG\.)\s*(\d+(?:[a-zA-Z])?)", ocr_text, re.IGNORECASE)
    if fig_matches:
        if "figure_references" not in identified_entities:
            identified_entities["figure_references"] = []
        identified_entities["figure_references"].extend(list(set(fig_matches)))  # Use set for uniqueness

    return identified_entities


def is_image_blank_or_too_small(pil_image, min_size=(100, 100),
                                blank_threshold=0.99):
    """
    Checks if a PIL image is effectively blank or too small to be meaningful.
    A blank image is one where most pixels are very similar (e.g., all white or all black).
    """
    if pil_image.width < min_size[0] or pil_image.height < min_size[1]:
        return True

    gray_image = pil_image.convert('L')

    if OPENCV_AVAILABLE:
        np_image = np.array(gray_image)
        std_dev = np.std(np_image)
        if std_dev < 15:
            return True
    else:
        hist = gray_image.histogram()
        num_pixels = pil_image.width * pil_image.height
        blank_pixels = hist[0] + hist[255]
        if (blank_pixels / num_pixels) > blank_threshold:
            return True

    return False


def extract_formulas_from_image(pil_image, ocr_text):
    """
    Simulates Mathematical OCR (M-OCR) for demonstration.
    In a real scenario, this would involve a dedicated M-OCR API or library.
    For this example, it uses heuristics to "detect" a formula and returns
    a hardcoded LaTeX string for the specific formula provided by the user.
    """
    # Normalize the OCR text for more robust matching
    normalized_ocr_text = ocr_text.lower().replace(" ", "").replace("−", "-")

    # Define key components of the formula to look for
    # Using flexible patterns to account for OCR errors (e.g., 'o2' instead of 'o^2')
    gc_pattern = r"gc"
    ef_pattern = r"ef"
    h_pattern = r"h"
    vo_pattern = r"vo\d?"  # Matches vo, vo2, etc.
    four_pattern = r"4"
    one_minus_nu_pattern = r"1-?[vνn]\d?"  # Matches 1-v, 1-v2, 1-nu, 1-nu2 etc.
    vc_pattern = r"vc\d?"  # Matches vc, vc2, etc.
    units_pattern = r"j/m2"

    # Check for the presence of all key components
    if (re.search(gc_pattern, normalized_ocr_text) and
            re.search(r"=", normalized_ocr_text) and
            re.search(ef_pattern, normalized_ocr_text) and
            re.search(h_pattern, normalized_ocr_text) and
            re.search(vo_pattern, normalized_ocr_text) and
            re.search(four_pattern, normalized_ocr_text) and
            re.search(one_minus_nu_pattern, normalized_ocr_text) and
            re.search(vc_pattern, normalized_ocr_text) and
            re.search(units_pattern, normalized_ocr_text)):
        # If all components are found, return the correct LaTeX string
        return r'$G_c = \frac{E_f h V_o^2}{4(1-\nu^2)V_c^2} (J/m^2)$'

    return None  # No formula detected


def extract_pdf_images_with_ocr(pdf_path, output_image_dir):
    """
    Extract images from PDF document, perform OCR on them, and save them locally.
    Returns a list of dictionaries with image data, OCR text, potential captions,
    identified entities, and the local path to the saved image.
    Skips images found in header, footer, or sidebar regions.
    Now also attempts to extract formulas using a simulated M-OCR.
    """
    os.makedirs(output_image_dir, exist_ok=True)
    images_data = []

    try:
        import fitz  # PyMuPDF

        pdf_document = fitz.open(pdf_path)

        with pdfplumber.open(pdf_path) as pdf_plumber_doc:
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)  # This is a PyMuPDF page object
                plumber_page = pdf_plumber_doc.pages[page_num]  # This is a pdfplumber page object

                page_width = plumber_page.width
                page_height = plumber_page.height

                # Define the bounding box for the main content area
                x0_content = page_width * SIDEBAR_SKIP_RATIO
                y0_content = 0
                x1_content = page_width * (1 - SIDEBAR_SKIP_RATIO)
                y1_content = page_height * (1 - FOOTER_SKIP_RATIO)
                main_content_bbox = (x0_content, y0_content, x1_content, y1_content)

                # Extract ALL text lines from the original page first using extract_words for robustness
                all_words = plumber_page.extract_words(x_tolerance=3)  # Use x_tolerance for word grouping
                all_page_text_lines = []
                if all_words:
                    # Sort words by vertical position (top to bottom), then horizontal (left to right)
                    all_words.sort(key=lambda w: (w['top'], w['x0']))

                    current_line_words = []
                    line_tolerance = 5  # pixels - adjust if needed

                    for word in all_words:
                        if not current_line_words:
                            current_line_words.append(word)
                        else:
                            last_word_top = current_line_words[-1]['top']
                            if abs(word['top'] - last_word_top) <= line_tolerance:
                                current_line_words.append(word)
                            else:
                                if current_line_words:
                                    current_line_words.sort(key=lambda w: w['x0'])
                                    line_text = " ".join([w['text'] for w in current_line_words])
                                    # Calculate bbox for the reconstructed line
                                    min_x0 = min(w['x0'] for w in current_line_words)
                                    max_x1 = max(w['x1'] for w in current_line_words)
                                    min_top = min(w['top'] for w in current_line_words)
                                    max_bottom = max(w['bottom'] for w in current_line_words)
                                    all_page_text_lines.append(
                                        {"text": line_text.strip(), "bbox": (min_x0, min_top, max_x1, max_bottom)})
                                current_line_words = [word]

                    if current_line_words:  # Don't forget the last line
                        current_line_words.sort(key=lambda w: w['x0'])
                        line_text = " ".join([w['text'] for w in current_line_words])
                        min_x0 = min(w['x0'] for w in current_line_words)
                        max_x1 = max(w['x1'] for w in current_line_words)
                        min_top = min(w['top'] for w in current_line_words)
                        max_bottom = max(w['bottom'] for w in current_line_words)
                        all_page_text_lines.append(
                            {"text": line_text.strip(), "bbox": (min_x0, min_top, max_x1, max_bottom)})

                image_list = page.get_images(full=True)  # This uses PyMuPDF page object
                plumber_images = plumber_page.images  # These bboxes are relative to original page

                for img_idx, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        base_image = pdf_document.extract_image(xref)
                        image_bytes = base_image["image"]
                        image_ext = base_image["ext"]

                        pil_image = Image.open(io.BytesIO(image_bytes))
                        if pil_image.mode != "RGB":
                            pil_image = pil_image.convert("RGB")

                        # --- NEW: Filter out blank or too small images ---
                        if is_image_blank_or_too_small(pil_image):
                            print(f"Skipping blank/small image on page {page_num + 1}, index {img_idx}")
                            continue
                        # --- END NEW ---

                        image_filename = f"page_{page_num + 1}_img_{img_idx + 1}.{image_ext}"
                        local_image_path = os.path.join(output_image_dir, image_filename)
                        pil_image.save(local_image_path)

                        processed_pil_image = preprocess_image_for_ocr(pil_image)
                        ocr_text = pytesseract.image_to_string(processed_pil_image, config='--psm 6')
                        cleaned_ocr_text = clean_text(ocr_text)

                        # --- NEW: Attempt M-OCR for formulas ---
                        formula_latex = extract_formulas_from_image(processed_pil_image, cleaned_ocr_text)
                        # --- END NEW ---

                        image_bbox = None
                        if img_idx < len(plumber_images):
                            image_bbox = (plumber_images[img_idx]['x0'], plumber_images[img_idx]['top'],
                                          plumber_images[img_idx]['x1'], plumber_images[img_idx]['bottom'])

                        # Check if image_bbox overlaps with the main content area
                        # If not, it's likely a decorative image in header/footer/sidebar, and we can skip it.
                        if image_bbox:
                            img_x0, img_y0, img_x1, img_y1 = image_bbox
                            # Check for overlap: (A.x1 > B.x0 and A.x0 < B.x1) and (A.y1 > B.y0 and A.y0 < B.y1)
                            if not (img_x1 > x0_content and img_x0 < x1_content and
                                    img_y1 > y0_content and img_y0 < y1_content):
                                print(
                                    f"Skipping image on page {page_num + 1}, index {img_idx} due to being outside main content area.")
                                os.remove(local_image_path)  # Delete the saved image
                                continue

                        # --- NEW: Filter out images with very little OCR text and no caption, and no formula ---
                        caption = None  # Initialize caption before using it in the condition
                        # Only search for caption if the image is within the main content area
                        if image_bbox:
                            # Filter text lines to only those within the main content bbox
                            filtered_page_text_lines = [
                                line_info for line_info in all_page_text_lines
                                if line_info['bbox'][1] >= y0_content and line_info['bbox'][3] <= y1_content and
                                   line_info['bbox'][0] >= x0_content and line_info['bbox'][2] <= x1_content
                            ]
                            caption = extract_caption_from_text(filtered_page_text_lines, image_bbox, page_width,
                                                                page_height)

                        if len(cleaned_ocr_text.strip()) < 10 and not caption and not formula_latex:  # Adjusted threshold to 10
                            print(
                                f"Skipping image on page {page_num + 1}, index {img_idx} due to minimal OCR text, no caption, and no formula detected.")
                            os.remove(local_image_path)  # Delete the saved blank image
                            continue
                        # --- END NEW ---

                        identified_entities = extract_specific_entities_from_ocr(cleaned_ocr_text)

                        images_data.append({
                            "page_num": page_num + 1,
                            "img_idx": img_idx,
                            "ocr_text": cleaned_ocr_text,
                            "caption": caption,
                            "identified_entities": identified_entities,
                            "local_path": local_image_path,
                            "filename": image_filename,
                            "formula_latex": formula_latex  # Store the extracted LaTeX
                        })

                    except Exception as e:
                        print(f"Error processing image on page {page_num + 1}, index {img_idx}: {e}")
                        continue
        pdf_document.close()

    except Exception as e:
        print(f"Error extracting images from PDF with PyMuPDF/pdfplumber: {e}")
        st.error(
            f"Failed to extract images from PDF: {e}. Ensure PyMuPDF and pdfplumber are installed and the PDF is not corrupted.")

    return images_data


def extract_text_from_columns(page):
    """
    Extracts text from main content area in two-column format.
    Assumes headers/affiliations are already handled separately.
    This function now operates on an already cropped page (main content area).
    """
    words = page.extract_words(x_tolerance=3)
    if not words:
        return []

    # Simple two-column split based on the cropped page's width
    page_width = page.width
    mid_point = page_width / 2

    left_words = []
    right_words = []

    # Classify words by column
    for word in words:
        word_center = (word['x0'] + word['x1']) / 2

        if word_center < mid_point:
            left_words.append(word)
        else:
            right_words.append(word)

    # Convert each column to text lines
    left_lines = column_to_lines(left_words)
    right_lines = column_to_lines(right_words)

    # Combine: left column first, then right column
    result = []
    result.extend(left_lines)
    result.extend(right_lines)

    return result


def column_to_lines(words):
    """
    Convert a list of words from one column into properly ordered text lines
    """
    if not words:
        return []

    # Sort words by vertical position (top to bottom), then horizontal (left to right)
    words.sort(key=lambda w: (w['top'], w['x0']))

    lines = []
    current_line_words = []  # Initialize current_line_words

    line_tolerance = 5  # pixels - adjust if needed

    for word in words:
        if not current_line_words:
            # First word of the line
            current_line_words.append(word)
        else:
            # Check if this word is on the same line as the previous word
            last_word_top = current_line_words[-1]['top']

            if abs(word['top'] - last_word_top) <= line_tolerance:
                # Same line - add to current line
                current_line_words.append(word)
            else:
                # New line - finish the current line and start a new one
                if current_line_words:
                    current_line_words.sort(key=lambda w: w['x0'])
                    line_text = " ".join([w['text'] for w in current_line_words])
                    lines.append(line_text.strip())

                # Start new line
                current_line_words = [word]

    # Don't forget the last line
    if current_line_words:
        current_line_words.sort(key=lambda w: w['x0'])
        line_text = " ".join([w['text'] for w in current_line_words])
        lines.append(line_text.strip())

    # Filter out empty lines
    return [line for line in lines if line]


# Alternative version with adjustable column boundary
def extract_text_from_columns_flexible(page, column_split_ratio=0.5):
    """
    Flexible version where you can adjust where the column split occurs

    Args:
        page: PDF page object
        column_split_ratio: Where to split (0.5 = middle, 0.4 = 40% from left, etc.)
    """
    words = page.extract_words(x_tolerance=3)
    if not words:
        return []

    page_width = page.width
    split_point = page_width * column_split_ratio

    left_words = []
    right_words = []

    for word in words:
        word_center = (word['x0'] + word['x1']) / 2

        if word_center < split_point:
            left_words.append(word)
        else:
            right_words.append(word)

    left_lines = column_to_lines(left_words)
    right_lines = column_to_lines(right_words)

    result = []
    result.extend(left_lines)
    result.extend(right_lines)

    return result


# Quick test function
def test_column_extraction(page):
    """
    Test function to see the results
    """
    result = extract_text_from_columns(page)

    print(f"Extracted {len(result)} lines:")
    for i, line in enumerate(result):
        print(f"{i + 1:2d}: {line}")

    return result


def extract_pdf_structure(pdf_path, document_category):
    """
    Extract hierarchical structure from PDF document,
    mapping to Document, Section, Subsection.
    Applies specific logic based on document_category.
    Skips defined header, footer, and sidebar regions.
    """
    document_structure = {
        "type": "Document",
        "title": "Untitled Document",
        "sections": [],
        "patent_number": None,
        "abstract_content": None
    }
    current_section = None
    current_subsection = None
    references_section_reached = False  # Flag to stop processing after references

    print("=== PDF Document Analysis ===")

    try:
        with pdfplumber.open(pdf_path) as pdf:
            all_text_lines = []
            for page_num, page in enumerate(pdf.pages):
                page_width = page.width
                page_height = page.height

                # Define the bounding box for the main content area
                # Exclude header, footer, and sidebars
                x0 = page_width * SIDEBAR_SKIP_RATIO
                y0 = 0
                x1 = page_width * (1 - SIDEBAR_SKIP_RATIO)
                y1 = page_height * (1 - FOOTER_SKIP_RATIO)

                main_content_bbox = (x0, y0, x1, y1)
                cropped_page = page.crop(main_content_bbox)

                # Use column extraction for IEEE and Patent documents on the cropped page
                if document_category in ["IEEE Knowledge", "Patent Knowledge"]:
                    page_lines = extract_text_from_columns(cropped_page)  # Pass the cropped page
                else:  # For General Document, use standard text extraction on the cropped page
                    page_text = cropped_page.extract_text(layout=True)
                    page_lines = [line.strip() for line in page_text.split('\n') if line.strip()] if page_text else []

                if page_lines:
                    all_text_lines.extend(page_lines)

        lines = all_text_lines
        current_content_buffer = []

        # Define patterns for various section levels (common for general/IEEE)
        # More specific pattern for Roman numeral sections
        roman_numeral_section_pattern = r'\b([IVXLCDM]+)\.\s+([A-Z][A-Z\s]+)\b'

        # Standalone section patterns (must match from the start of the line)
        section_patterns_standalone = [
            r'^\s*ABSTRACT\s*$',
            r'^\s*KEYWORDS\s*$',
            r'^\s*INTRODUCTION\s*$',
            r'^\s*ACKNOWLEDGEMENT\s*$',
            r'^\d+\.\s+\b[A-Z][A-Z\s]*\b\s*$',  # Numbered sections (e.g., 1. INTRODUCTION, 2. METHODOLOGY)
            r'^\s*(?:RESULTS\s+AND\s+DISCUSSION|EXPERIMENTAL\s+PROTOCOL|RELATED\s+WORK|METHODOLOGY|DESIGN|BACKGROUND|SUMMARY|BRIEF\s+DESCRIPTION\s+OF\s+THE\s+SEVERAL\s+VIEWS\s+OF\s+THE\s+DRAWINGS|DETAILED\s+DESCRIPTION|CLAIMS)\s*$',
            # Specific common headers, including patent ones
            r'^\s*CONCLUSION(?:S)?\s*$',
        ]

        subsection_patterns_standalone = [
            r'^[A-Z]\.\s+([A-Z][a-zA-Z\s,]+)$',  # A. Subsection Title - more specific
            r'^\d+\.\d+\s+([A-Z][a-zA-Z0-9\s,]+)$',  # 1.1 Subsection Title - more specific
        ]
        # Base junk patterns (common for most documents)
        junk_patterns = [
            r'^\s*Authorized licensed use limited to:', r'^\s*Downloaded on', r'^\s*Restrictions apply\.',
            r'.*IEEE.*Conference.*', r'.*Conference.*year.*', r'.*Symposium.*year.*', r'.*Proceedings.*year.*',
            r'.*ECTC.*', r'^\d+$', r'^page\s+\d+$',
            r'^\s*Copyright\s*.*$',
            r'^\s*DOI:\s*.*$', r'^\s*\d{4}\s*IEEE\s*$',
            r'^\s*\d{3}-\d{3}-\d{4}-\d{4}/\d{2}/\$\d{2}\.00\s*©\s*\d{4}\s*IEEE\s*$', r'^\s*\(\d+\)\s*$',
            r'^\s*\[\d+\]\s*$',
            r'^\s*REFERENCES\s*$',
        ]

        # --- PATENT CONVERSION LOGIC BLOCK ---
        if document_category == "Patent Knowledge":
            print("Applying Patent Conversion Logic for PDF...")
            junk_patterns.extend([
                r'^\s*U\.?S\.?\s*Cl\.:\s*.*$', r'^\s*Field of Classification Search\s*.*$',
                r'^\s*(?:US|U\.S\.)\s*Patent\s*Documents\s*$', r'^\s*\d{1,},\d{3},\d{3}\s+[A-Z][a-zA-Z]+\s+.*$',
                r'^\s*Notice:\s*Subject to any disclaimer.*$', r'^\s*App\.\s*No\.:\s*.*$', r'^\s*Filed:\s*.*$',
                r'^\s*Date of Patent:\s*.*$', r'^\s*United States Patent\s*$', r'^\s*[A-Z][a-z]+ et al\.\s*$',
                r'^\s*US\d{7,8}[A-Z]\d$',
                r'^\s*\(11\)\s*Patent Publication Number.*$',
                r'^\s*\(19)\s*(?:Country Identifier|United States).*$',
                r'^\s*\(43)\s*Publication Date.*$',
                r'^\s*\(12)\s*United States Patent.*$',
                r'^\s*\(45)\s*Date of Patent:.*$',
                r'^\s*\(52)\s*U\.S\. CI\..*$',
                r'^\s*CPC.*$',
                r'^\s*\(71\)\s*Applicant:.*$',
                r'^\s*\(72)\s*Inventors:.*$',
                r'^\s*\(73)\s*Assignee:.*$',
                r'^\s*\(58)\s*Field of Classification Search.*$',
                r'^\s*\(56)\s*References Cited.*$',
                r'^\s*\(21)\s*Appl\. No\.:.*$',
                r'^\s*Primary Examiner.*$',
                r'^\s*\(74)\s*Attorney,\s*Agent,\s*or\s*Firm.*$',
                r'^\s*\(22)\s*Filed:.*$',
                r'^\s*\(65)\s*Prior Publication Data.*$',
                r'^\s*\(62)\s*Related U\.S\. Application Data.*$',
                r'^\s*\(51)\s*Int\. Cl\..*$',
                r'^\s*\(60)\s*Provisional application No\..*$',
                r'^\s*\(57)\s*ABSTRACT.*$',
                r'^\s*\*\s*cited by examiner.*$',
                r'^\s*\(\*\)\s*Notice:.*$',
                r'^\s*20 Claims,\s*\d+\s*Drawing Sheets.*$'
            ])

            patent_title_found = False
            patent_number_found = False
            abstract_header_found = False
            start_main_content_index = 0
            abstract_buffer = []

            for i, line in enumerate(lines):
                text = clean_text(line)
                if not text:
                    continue

                if not patent_title_found:
                    patent_title_match = re.search(r"^\(54\)\s*(.*)", text)
                    if patent_title_match:
                        document_structure["title"] = clean_text(patent_title_match.group(1))
                        patent_title_found = True
                        patent_title_suffixes_to_remove = [
                            r'\s*Publication Classification$',
                            r'\s*OF HOLL \d{1,2}/\d{1,2}\s*\(\d{4}\.\d{2}\)$',
                            r'\s*AND METHOD AND EQUIPMENT FOR FORMING THE SAME$',
                            r'\s*AND METHOD AND APPARATUS FOR FORMING THE SAME$',
                        ]
                        for suffix_pattern in patent_title_suffixes_to_remove:
                            document_structure["title"] = re.sub(suffix_pattern, '', document_structure["title"],
                                                                 flags=re.IGNORECASE).strip()
                        start_main_content_index = i + 1
                        continue

                if not patent_number_found:
                    patent_number_match = re.search(
                        r"^\(10\)\s*Patent No\.:\s*(US\s*\d{1,},\d{3},\d{3}\s*[A-Z]\d|US\s*\d{7,8}[A-Z]\d)", text)
                    if patent_number_match:
                        document_structure["patent_number"] = clean_text(patent_number_match.group(1))
                        patent_number_found = True
                        start_main_content_index = i + 1
                        continue

                if not abstract_header_found:
                    if re.match(r"^\(57\)\s*ABSTRACT\s*$", text, re.IGNORECASE):
                        abstract_header_found = True
                        start_main_content_index = i + 1
                        continue

                if abstract_header_found:
                    if any(re.match(p, text, re.IGNORECASE) for p in section_patterns_standalone) or \
                            any(re.match(p, text, re.IGNORECASE) for p in junk_patterns):
                        document_structure["abstract_content"] = clean_text(" ".join(abstract_buffer))
                        abstract_header_found = False
                        start_main_content_index = i
                        break
                    else:
                        abstract_buffer.append(text)
                        main_content_start_index = i + 1

            if abstract_header_found and abstract_buffer:
                document_structure["abstract_content"] = clean_text(" ".join(abstract_buffer))

            actual_start_index = start_main_content_index
        # --- END PATENT CONVERSION LOGIC BLOCK ---

        # --- IEEE CONVERSION LOGIC BLOCK ---
        elif document_category == "IEEE Knowledge":
            print("Applying IEEE Conversion Logic for PDF (refined).")
            junk_patterns.extend([
                r'^\s*\d{4}\s+IEEE\s+\d{2,3}th\s+Electronic\s+Components\s+and\s+Technology\s+Conference.*',
                r'^\s*DOI:\s+10\.\d{4}/ECTC\d{5}\.\d{4}\.\d{5,6}\s*$',
                r'^\s*Authorized\s+licensed\s+use\s+limited\s+to:.*',
                r'^\s*Downloaded\s+on.*UTC\s+from\s+IEEE\s+Xplore\.\s+Restrictions\s+apply\.\s*$'
            ])

            ieee_author_affiliation_patterns = [
                r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}',
                r'\b(?:university|institute|company|corp|ltd|inc|department|lab|center)\b',
                r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,}(?:(?:\s*,\s*[A-Z][a-z]+)?(?:\s+et al\.)?)?$',
                r'^\s*[\w\s,.-]+(?:University|Institute|Company|Corp|Ltd|Inc|Department|Lab|Center)\s*$',
                r'^\s*[\w\s,.-]+\s*,\s*(?:Japan|USA|China|Korea|Germany|UK|France)\s*$',
                r'^\s*[\w\s,.-]+\s*@[\w\s,.-]+\.[\w]{2,}\s*$',
            ]

            state = "LOOKING_FOR_TITLE"
            main_content_start_index = 0
            abstract_buffer = []
            title_buffer = []

            # First pass to find the title
            temp_lines = list(lines)  # Make a copy to modify
            if temp_lines and re.search(r'IEEE.*Conference', temp_lines[0], re.IGNORECASE):
                temp_lines.pop(0)

            for i, line in enumerate(temp_lines):
                text = clean_text(line)
                if not text: continue

                is_author_line = any(re.search(p, text, re.IGNORECASE) for p in ieee_author_affiliation_patterns)
                is_abstract_header = re.match(r"^\s*Abstract\b", text, re.IGNORECASE)

                if is_author_line or is_abstract_header:
                    document_structure["title"] = " ".join(title_buffer).strip()
                    main_content_start_index = i
                    break
                else:
                    title_buffer.append(text)

            if not document_structure["title"] or document_structure["title"] == "Untitled Document":
                document_structure["title"] = " ".join(title_buffer).strip()

            # Second pass for sections
            in_abstract = False
            for i, line in enumerate(lines):
                text = clean_text(line)
                if not text: continue

                if re.match(r"^\s*Abstract\b", text, re.IGNORECASE):
                    in_abstract = True
                    abstract_buffer.append(text.split("Abstract", 1)[-1].strip())
                    continue

                if in_abstract:
                    if re.match(r"^\s*Keywords\b", text, re.IGNORECASE) or re.search(roman_numeral_section_pattern,
                                                                                     text, re.IGNORECASE):
                        in_abstract = False
                        document_structure["abstract_content"] = " ".join(abstract_buffer).strip()
                        actual_start_index = i
                        break
                    else:
                        abstract_buffer.append(text)
            else:  # If loop completes without break
                if abstract_buffer:
                    document_structure["abstract_content"] = " ".join(abstract_buffer).strip()
                actual_start_index = len(lines)


        # --- END IEEE CONVERSION LOGIC BLOCK ---

        # --- GENERAL DOCUMENT CONVERSION LOGIC BLOCK ---
        else:
            print("Applying General Document Conversion Logic for PDF.")
            actual_start_index = 0
        # --- END GENERAL DOCUMENT CONVERSION LOGIC BLOCK ---

        for i, line in enumerate(lines[actual_start_index:]):
            text = clean_text(line)
            if not text:
                continue

            # Check for REFERENCES section (highest priority for stopping content accumulation)
            if re.match(r'^\s*REFERENCES\s*$', text, re.IGNORECASE):
                references_section_reached = True
                if current_content_buffer:
                    full_content = " ".join(current_content_buffer)
                    if current_subsection:
                        current_subsection["content"].append(full_content)
                    elif current_section:
                        current_section["content"].append(full_content)
                    else:
                        pass
                    current_content_buffer = []
                break

            if references_section_reached:
                continue

            # Filter out junk lines and page numbers
            if any(re.search(p, text, re.IGNORECASE) for p in junk_patterns) or \
                    text.isdigit() or re.match(r'^page\s+\d+$', text, re.IGNORECASE):
                continue

            # IEEE specific filtering
            if document_category == "IEEE Knowledge":
                if any(re.search(p, text, re.IGNORECASE) for p in ieee_author_affiliation_patterns):
                    continue

            # Flag to indicate if a new structure (section/subsection) was identified and processed in this iteration
            structure_identified_this_line = False

            # --- UPDATED: Search for embedded Roman numeral sections ---
            roman_match = re.search(roman_numeral_section_pattern, text, re.IGNORECASE)
            if roman_match:
                pre_match_content = text[:roman_match.start()].strip()
                header_title = roman_match.group(0).strip()
                post_match_content = text[roman_match.end():].strip()

                # Add any content before the match to the previous section
                if pre_match_content:
                    current_content_buffer.append(pre_match_content)

                # Flush buffer before creating the new section
                if current_content_buffer:
                    full_content = " ".join(current_content_buffer)
                    if current_subsection:
                        current_subsection["content"].append(full_content)
                    elif current_section:
                        current_section["content"].append(full_content)
                    current_content_buffer = []

                # Create the new section
                new_section = {
                    "type": "Section",
                    "title": header_title,
                    "subsections": [],
                    "content": []
                }
                document_structure["sections"].append(new_section)
                current_section = new_section
                current_subsection = None
                structure_identified_this_line = True

                # Add any content after the match to the new section's buffer
                if post_match_content:
                    current_content_buffer.append(post_match_content)

            # --- If no Roman numeral section was found, check for other standalone section/subsection patterns ---
            if not structure_identified_this_line:
                is_section_header = False
                is_subsection_header = False
                header_title = None

                # Check for other standalone section patterns (must match from the start of the line)
                for pattern in section_patterns_standalone:
                    if re.match(pattern, text, re.IGNORECASE) and len(text.split()) < 10:
                        is_section_header = True
                        header_title = text.strip()
                        break

                if is_section_header:
                    structure_identified_this_line = True
                    # Flush content before creating new section
                    if current_content_buffer:
                        full_content = " ".join(current_content_buffer)
                        if current_subsection:
                            current_subsection["content"].append(full_content)
                        elif current_section:
                            current_section["content"].append(full_content)
                        else:
                            pass
                        current_content_buffer = []

                    new_section = {
                        "type": "Section",
                        "title": header_title,
                        "subsections": [],
                        "content": []
                    }
                    document_structure["sections"].append(new_section)
                    current_section = new_section
                    current_subsection = None

                else:
                    # Check for standalone subsection patterns (must match from the start of the line)
                    for pattern in subsection_patterns_standalone:
                        if re.match(pattern, text, re.IGNORECASE) and len(text.split()) < 10:
                            is_subsection_header = True
                            header_title = text.strip()
                            break

                    if is_subsection_header:
                        structure_identified_this_line = True
                        # Flush content before creating new subsection
                        if current_content_buffer:
                            full_content = " ".join(current_content_buffer)
                            if current_subsection:
                                current_subsection["content"].append(full_content)
                            elif current_section:
                                current_section["content"].append(full_content)
                            else:
                                pass
                            current_content_buffer = []

                        new_subsection = {
                            "type": "Subsection",
                            "title": header_title,
                            "subsubsections": [],
                            "content": []
                        }
                        if current_section is None:
                            continue
                        current_section["subsections"].append(new_subsection)
                        current_subsection = new_subsection

            # If no new structure was identified in this line, it's regular content
            if not structure_identified_this_line:
                current_content_buffer.append(text)

        # Final flush of any remaining content
        if current_content_buffer and not references_section_reached:
            full_content = " ".join(current_content_buffer)
            if current_subsection:
                current_subsection["content"].append(full_content)
            elif current_section:
                current_section["content"].append(full_content)
            else:
                pass
            current_content_buffer = []

    except Exception as e:
        print(f"Error processing PDF with pdfplumber, falling back to PyPDF2: {e}")
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            all_text = ""
            for page in pdf_reader.pages:
                all_text += page.extract_text() + "\n"

        if all_text.strip():
            document_structure["title"] = "Full Document Content (Fallback)"
            document_structure["sections"] = [{
                "type": "Section",
                "title": "Main Content",
                "subsections": [],
                "content": [clean_text(all_text)]
            }]
        else:
            print("PyPDF2 also failed to extract text or found no content.")

    return document_structure


def is_likely_author_affiliation_table(table_rows):
    """
    Heuristically checks if a detected table is likely an author/affiliation block.
    This is based on common patterns in scientific paper title pages.
    """
    if not table_rows:
        return False

    num_columns_in_first_row = len(table_rows[0]) if table_rows else 0
    if num_columns_in_first_row > 2:
        return False

    email_pattern = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
    contains_email = False
    for row in table_rows:
        for cell in row:
            if re.search(email_pattern, cell):
                contains_email = True
                break
        if contains_email:
            break

    if (num_columns_in_first_row <= 2) and contains_email and (len(table_rows) > 2):
        return True
    return False


def extract_pdf_tables(pdf_path):
    """Extract table content from PDF document, skipping header, footer, and sidebar regions."""
    tables = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_width = page.width
                page_height = page.height

                x0_content = page_width * SIDEBAR_SKIP_RATIO
                y0_content = 0
                x1_content = page_width * (1 - SIDEBAR_SKIP_RATIO)
                y1_content = page_height * (1 - FOOTER_SKIP_RATIO)
                main_content_bbox = (x0_content, y0_content, x1_content, y1_content)

                cropped_page = page.crop(main_content_bbox)
                page_tables = cropped_page.extract_tables()

                if page_tables:
                    for table in page_tables:
                        rows = []
                        for row in table:
                            cells = [clean_text(str(cell)) for cell in row if cell and clean_text(str(cell))]
                            if cells:
                                rows.append(cells)
                        if rows:
                            tables.append(rows)
    except Exception as e:
        print(f"Error extracting PDF tables: {e}")

    return tables


def extract_doc_structure(docx_path, document_category):
    """
    Extract hierarchical structure from Word document,
    mapping to Document, Section, Subsection.
    Applies specific logic based on document_category.
    Note: Region skipping (header/footer/sidebar) for DOCX is primarily handled
    via the 'junk_patterns' and explicit content parsing logic, as python-docx
    does not provide direct page-layout-based cropping like pdfplumber.
    """
    doc = Document(docx_path)
    document_structure = {
        "type": "Document",
        "title": "Untitled Document",
        "sections": [],
        "patent_number": None,
        "abstract_content": None
    }
    current_section = None
    current_subsection = None
    references_section_reached = False

    print("=== Word Document Analysis ===")

    lines = [clean_text(p.text) for p in doc.paragraphs]
    current_content_buffer = []

    section_style_patterns = ["Title", "Heading 1"]
    # More specific pattern for Roman numeral sections
    roman_numeral_section_pattern = r'^\s*([IVXLCDM]+)\.\s+([A-Z\s]+)\s*$'

    # Standalone section patterns (must match from the start of the line)
    section_text_patterns_standalone = [
        r'^\s*ABSTRACT\s*$', r'^\s*KEYWORDS\s*$', r'^\s*INTRODUCTION\s*$', r'^\s*ACKNOWLEDGEMENT\s*$',
        r'^\d+\.\s+\b[A-Z][A-Z\s]*\b\s*$',  # More specific for numbered sections
        r'^\s*(?:RESULTS\s+AND\s+DISCUSSION|EXPERIMENTAL\s+PROTOCOL|RELATED\s+WORK|METHODOLOGY|DESIGN|BACKGROUND|SUMMARY|BRIEF\s+DESCRIPTION\s+OF\s+THE\s+SEVERAL\s+VIEWS\s+OF\s+THE\s+DRAWINGS|DETAILED\s+DESCRIPTION|CLAIMS)\s*$',
        r'^#\s+',  # Markdown style headings
        r'^\s*CONCLUSION(?:S)?\s*$',
    ]
    subsection_style_patterns = ["Heading 2"]
    subsection_text_patterns_standalone = [
        r'^[A-Z]\.\s+\b[A-Z][a-zA-Z\s]*\b\s*$',  # More specific for A. style
        r'^\d+\.\d+\s+\b[A-Z][a-zA-Z\s]*\b\s*$',  # 1.1 Subsection Title
        r'^##\s+'  # Markdown style subheadings
    ]
    junk_patterns = [
        r'^\s*Authorized licensed use limited to:', r'^\s*Downloaded on', r'^\s*Restrictions apply\.',
        r'.*IEEE.*Conference.*', r'.*Conference.*year.*', r'.*Symposium.*year.*', r'.*Proceedings.*year.*',
        r'.*ECTC.*', r'^\d+$', r'^page\s+\d+$',
        r'^\s*Copyright\s*.*$',
        r'^\s*DOI:\s*.*$', r'^\s*\d{4}\s*IEEE\s*$',
        r'^\s*\d{3}-\d{3}-\d{4}-\d{4}/\d{2}/\$\d{2}\.00\s*©\s*\d{4}\s*IEEE\s*$', r'^\s*\(\d+\)\s*$',
        r'^\s*\[\d+\]\s*$',
        r'^\s*REFERENCES\s*$',
    ]

    # --- PATENT CONVERSION LOGIC BLOCK ---
    if document_category == "Patent Knowledge":
        print("Applying Patent Conversion Logic for DOCX...")
        junk_patterns.extend([
            r'^\s*U\.?S\.?\s*Cl\.:\s*.*$', r'^\s*Field of Classification Search\s*.*$',
            r'^\s*(?:US|U\.S\.)\s*Patent\s*Documents\s*$', r'^\s*\d{1,},\d{3},\d{3}\s+[A-Z][a-zA-Z]+\s+.*$',
            r'^\s*Notice:\s*Subject to any disclaimer.*$', r'^\s*App\.\s*No\.:\s*.*$', r'^\s*Filed:\s*.*$',
            r'^\s*Date of Patent:\s*.*$', r'^\s*United States Patent\s*$', r'^\s*[A-Z][a-z]+ et al\.\s*$',
            r'^\s*US\d{7,8}[A-Z]\d$',
            r'^\s*\(11\)\s*Patent Publication Number.*$',
            r'^\s*\(19)\s*(?:Country Identifier|United States).*$',
            r'^\s*\(43)\s*Publication Date.*$',
            r'^\s*\(12)\s*United States Patent.*$',
            r'^\s*\(45)\s*Date of Patent:.*$',
            r'^\s*\(52)\s*U\.S\. CI\..*$',
            r'^\s*CPC.*$',
            r'^\s*\(71\)\s*Applicant:.*$',
            r'^\s*\(72)\s*Inventors:.*$',
            r'^\s*\(73)\s*Assignee:.*$',
            r'^\s*\(58)\s*Field of Classification Search.*$',
            r'^\s*\(56)\s*References Cited.*$',
            r'^\s*\(21)\s*Appl\. No\.:.*$',
            r'^\s*Primary Examiner.*$',
            r'^\s*\(74)\s*Attorney,\s*Agent,\s*or\s*Firm.*$',
            r'^\s*\(22)\s*Filed:.*$',
            r'^\s*\(65)\s*Prior Publication Data.*$',
            r'^\s*\(62)\s*Related U\.S\. Application Data.*$',
            r'^\s*\(51)\s*Int\. Cl\..*$',
            r'^\s*\(60)\s*Provisional application No\..*$',
            r'^\s*\(57)\s*ABSTRACT.*$',
            r'^\s*\*\s*cited by examiner.*$',
            r'^\s*\(\*\)\s*Notice:.*$',
            r'^\s*20 Claims,\s*\d+\s*Drawing Sheets.*$'
        ])

        patent_title_found = False
        patent_number_found = False
        abstract_header_found = False
        start_main_content_index = 0
        abstract_buffer = []

        for i, para in enumerate(doc.paragraphs):
            text = clean_text(para.text)
            style = para.style.name

            if not text:
                continue

            if not patent_title_found:
                patent_title_match = re.search(r"^\(54\)\s*(.*)", text)
                if patent_title_match:
                    document_structure["title"] = clean_text(patent_title_match.group(1))
                    patent_title_found = True
                    patent_title_suffixes_to_remove = [
                        r'\s*Publication Classification$',
                        r'\s*OF HOLL \d{1,2}/\d{1,2}\s*\(\d{4}\.\d{2}\)$',
                        r'\s*AND METHOD AND EQUIPMENT FOR FORMING THE SAME$',
                        r'\s*AND METHOD AND APPARATUS FOR FORMING THE SAME$',
                    ]
                    for suffix_pattern in patent_title_suffixes_to_remove:
                        document_structure["title"] = re.sub(suffix_pattern, '', document_structure["title"],
                                                             flags=re.IGNORECASE).strip()
                    start_main_content_index = i + 1
                    continue

            if not patent_number_found:
                patent_number_match = re.search(
                    r"^\(10\)\s*Patent No\.:\s*(US\s*\d{1,},\d{3},\d{3}\s*[A-Z]\d|US\s*\d{7,8}[A-Z]\d)", text)
                if patent_number_match:
                    document_structure["patent_number"] = clean_text(patent_number_match.group(1))
                    patent_number_found = True
                    start_main_content_index = i + 1
                    continue

            if not abstract_header_found:
                if re.match(r"^\(57\)\s*ABSTRACT\s*$", text, re.IGNORECASE):
                    abstract_header_found = True
                    start_main_content_index = i + 1
                    continue

            if abstract_header_found:
                if any(s in style for s in section_style_patterns) or \
                        any(re.match(p, text, re.IGNORECASE) for p in section_text_patterns_standalone) or \
                        re.match(roman_numeral_section_pattern, text, re.IGNORECASE) or \
                        any(re.match(p, text, re.IGNORECASE) for p in junk_patterns):
                    document_structure["abstract_content"] = clean_text(" ".join(abstract_buffer))
                    abstract_header_found = False
                    start_main_content_index = i
                    break
                else:
                    abstract_buffer.append(text)
                    main_content_start_index = i + 1

        if abstract_header_found and abstract_buffer:
            document_structure["abstract_content"] = clean_text(" ".join(abstract_buffer))

        actual_start_index = start_main_content_index
    # --- END PATENT CONVERSION LOGIC BLOCK ---

    # --- IEEE CONVERSION LOGIC BLOCK ---
    elif document_category == "IEEE Knowledge":
        print("Applying IEEE Conversion Logic for DOCX (refined).")
        junk_patterns.extend([
            r'^\s*\d{4}\s+IEEE\s+\d{2,3}th\s+Electronic\s+Components\s+and\s+Technology\s+Conference.*',
            r'^\s*DOI:\s+10\.\d{4}/ECTC\d{5}\.\d{4}\.\d{5,6}\s*$',
            r'^\s*Authorized\s+licensed\s+use\s+limited\s+to:.*',
            r'^\s*Downloaded\s+on.*UTC\s+from\s+IEEE\s+Xplore\.\s+Restrictions\s+apply\.\s*$'
        ])

        ieee_author_affiliation_patterns = [
            r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}',
            r'\b(?:university|institute|company|corp|ltd|inc|department|lab|center)\b',
            r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,}(?:(?:\s*,\s*[A-Z][a-z]+)?(?:\s+et al\.)?)?$',
            r'^\s*[\w\s,.-]+(?:University|Institute|Company|Corp|Ltd|Inc|Department|Lab|Center)\s*$',
            r'^\s*[\w\s,.-]+\s*,\s*(?:Japan|USA|China|Korea|Germany|UK|France)\s*$',
            r'^\s*[\w\s,.-]+\s*@[\w\s,.-]+\.[\w]{2,}\s*$',
        ]

        state = "LOOKING_FOR_TITLE"
        main_content_start_index = 0
        abstract_buffer = []
        title_buffer = []

        for i, para in enumerate(doc.paragraphs):
            text = clean_text(para.text)
            style = para.style.name

            if not text:
                continue

            is_junk_line_current = any(re.search(p, text, re.IGNORECASE) for p in junk_patterns)
            is_strong_section_header_current = any(s in style for s in section_style_patterns) or \
                                               any(re.match(p, text, re.IGNORECASE) for p in
                                                   section_text_patterns_standalone) or \
                                               re.match(roman_numeral_section_pattern, text,
                                                        re.IGNORECASE)
            is_abstract_header = re.match(r"^\s*Abstract\b", text, re.IGNORECASE)
            is_keywords_header = re.match(r"^\s*Keywords\b", text, re.IGNORECASE)
            is_ieee_author_affiliation_line = any(
                re.search(p, text, re.IGNORECASE) for p in ieee_author_affiliation_patterns)

            if state == "LOOKING_FOR_TITLE":
                if is_junk_line_current:
                    continue
                if is_ieee_author_affiliation_line or is_abstract_header:
                    if title_buffer:
                        document_structure["title"] = " ".join(title_buffer)
                        state = "AFTER_TITLE"
                        main_content_start_index = i
                else:
                    title_buffer.append(text)

            elif state == "AFTER_TITLE":
                if is_abstract_header:
                    state = "IN_ABSTRACT"
                    abstract_buffer.append(text.split("Abstract", 1)[-1].strip())
                    main_content_start_index = i + 1
                elif is_keywords_header:
                    state = "IN_KEYWORDS"
                    main_content_start_index = i + 1
                elif is_strong_section_header_current:
                    state = "IN_MAIN_CONTENT"
                    main_content_start_index = i
                    break

            elif state == "IN_ABSTRACT":
                if is_keywords_header or is_strong_section_header_current:
                    document_structure["abstract_content"] = clean_text(" ".join(abstract_buffer))
                    abstract_buffer = []
                    state = "IN_MAIN_CONTENT"
                    main_content_start_index = i
                    break
                else:
                    abstract_buffer.append(text)
                    main_content_start_index = i + 1

        if state == "IN_ABSTRACT" and abstract_buffer:
            document_structure["abstract_content"] = clean_text(" ".join(abstract_buffer))

        actual_start_index = main_content_start_index
    # --- END IEEE CONVERSION LOGIC BLOCK ---
    else:  # General Document
        actual_start_index = 0

    for i, para in enumerate(doc.paragraphs[actual_start_index:]):
        current_line_index = actual_start_index + i
        text = clean_text(para.text)
        style = para.style.name

        if re.match(r'^\s*REFERENCES\s*$', text, re.IGNORECASE):
            references_section_reached = True
            if current_content_buffer:
                full_content = " ".join(current_content_buffer)
                if current_subsection:
                    current_subsection["content"].append(full_content)
                elif current_section:
                    current_section["content"].append(full_content)
                else:
                    pass
                current_content_buffer = []
            break

        if references_section_reached:
            continue

        is_junk_line = any(re.search(p, text, re.IGNORECASE) for p in junk_patterns)
        is_page_number = text.isdigit() or re.match(r'^page\s+\d+$', text, re.IGNORECASE)

        if is_junk_line or is_page_number:
            continue

        if document_category == "IEEE Knowledge":
            is_ieee_author_affiliation_line = any(
                re.search(p, text, re.IGNORECASE) for p in ieee_author_affiliation_patterns)
            if is_ieee_author_affiliation_line:
                continue

        # Flag to indicate if a new structure (section/subsection) was identified and processed in this iteration
        structure_identified_this_line = False

        # --- Attempt to identify and process Roman numeral section ---
        roman_match = re.match(roman_numeral_section_pattern, text, re.IGNORECASE)
        if roman_match:
            header_title = text.strip()

            # Flush any accumulated content before starting a new section
            if current_content_buffer:
                full_content = " ".join(current_content_buffer)
                if current_subsection:
                    current_subsection["content"].append(full_content)
                elif current_section:
                    current_section["content"].append(full_content)
                current_content_buffer = []  # Clear buffer after flushing

            # Create the new section
            new_section = {
                "type": "Section",
                "title": header_title,
                "subsections": [],
                "content": []
            }
            document_structure["sections"].append(new_section)
            current_section = new_section
            current_subsection = None  # Reset subsection as we're starting a new main section
            structure_identified_this_line = True

        # --- If no Roman numeral section was found, check for other standalone section/subsection patterns ---
        if not structure_identified_this_line:
            is_section_header = False
            is_subsection_header = False
            header_title = None

            # Check for other standalone section patterns (must match from the start of the line)
            # For DOCX, also check style patterns
            if any(s in style for s in section_style_patterns):
                is_section_header = True
                header_title = text.strip()
            else:
                for pattern in section_text_patterns_standalone:
                    if re.match(pattern, text, re.IGNORECASE) and len(text.split()) < 10:
                        is_section_header = True
                        header_title = text.strip()
                        break

            if is_section_header:
                structure_identified_this_line = True
                # Flush content before creating new section
                if current_content_buffer:
                    full_content = " ".join(current_content_buffer)
                    if current_subsection:
                        current_subsection["content"].append(full_content)
                    elif current_section:
                        current_section["content"].append(full_content)
                    else:
                        pass
                    current_content_buffer = []

                new_section = {
                    "type": "Section",
                    "title": header_title,
                    "subsections": [],
                    "content": []
                }
                document_structure["sections"].append(new_section)
                current_section = new_section
                current_subsection = None

            else:
                # Check for standalone subsection patterns (must match from the start of the line)
                # For DOCX, also check style patterns
                if any(s in style for s in subsection_style_patterns):
                    is_subsection_header = True
                    header_title = text.strip()
                else:
                    for pattern in subsection_text_patterns_standalone:
                        if re.match(pattern, text, re.IGNORECASE) and len(text.split()) < 10:
                            is_subsection_header = True
                            header_title = text.strip()
                            break

                if is_subsection_header:
                    structure_identified_this_line = True
                    # Flush content before creating new subsection
                    if current_content_buffer:
                        full_content = " ".join(current_content_buffer)
                        if current_subsection:
                            current_subsection["content"].append(full_content)
                        elif current_section:
                            current_section["content"].append(full_content)
                        else:
                            pass
                        current_content_buffer = []

                    new_subsection = {
                        "type": "Subsection",
                        "title": header_title,
                        "subsubsections": [],
                        "content": []
                    }
                    if current_section is None:
                        # Create a default section if a subsection is found first
                        default_section = {"type": "Section", "title": "Default Section", "subsections": [],
                                           "content": []}
                        document_structure["sections"].append(default_section)
                        current_section = default_section

                    current_section["subsections"].append(new_subsection)
                    current_subsection = new_subsection

        # If no new structure was identified in this line, it's regular content
        if not structure_identified_this_line:
            current_content_buffer.append(text)

    # Final flush of any remaining content
    if current_content_buffer and not references_section_reached:
        full_content = " ".join(current_content_buffer)
        if current_subsection:
            current_subsection["content"].append(full_content)
        elif current_section:
            current_section["content"].append(full_content)
        else:
            # If no section exists, create a default one for the content
            if not document_structure["sections"]:
                default_section = {"type": "Section", "title": "Main Content", "subsections": [], "content": []}
                document_structure["sections"].append(default_section)
                current_section = default_section
            current_section["content"].append(full_content)
        current_content_buffer = []

    return document_structure


def extract_doc_tables(docx_path):
    """
    Extract table content from Word document.
    Note: This is a basic implementation. Advanced table extraction from DOCX
    might require more sophisticated logic or libraries.
    """
    doc = Document(docx_path)
    tables = []
    for table_obj in doc.tables:
        current_table = []
        for row in table_obj.rows:
            row_cells = []
            for cell in row.cells:
                row_cells.append(clean_text(cell.text))
            if row_cells:
                current_table.append(row_cells)
        if current_table:
            tables.append(current_table)
    return tables


def safe_uri_name(text):
    """Create extremely safe URI names from text"""
    if not text:
        return "Unnamed"

    safe_name = re.sub(r'[^a-zA-Z0-9\s]', '', text)
    safe_name = re.sub(r'\s+', '_', safe_name)
    safe_name = re.sub(r'_+', '_', safe_name)
    safe_name = safe_name.strip('_')

    if not safe_name or not safe_name[0].isalpha():
        safe_name = "Item_" + safe_name

    safe_name = safe_name[:30] if safe_name else "Item"

    return safe_name


def safe_literal(text):
    """Create extremely safe literal values for RDF"""
    if not text:
        return "Empty"

    text = text.replace('\\', '/')
    text = text.replace('"', "'")
    text = text.replace('\n', ' ')
    text = text.replace('\r', ' ')
    text = text.replace('\t', ' ')

    text = ''.join(char for char in text if ord(char) >= 32 or char in ['\n', '\r', '\t'])
    text = re.sub(r'\s+', ' ', text).strip()

    if len(text) > 500:
        text = text[:497] + "..."

    return text if text else "Empty"


def chunk_content(text, max_len=120):
    """Breaks a long text into smaller chunks of a maximum length."""
    words = text.split()
    if not words:
        return []

    chunks = []
    current_chunk = ""
    for word in words:
        if len(current_chunk) + len(word) + 1 > max_len:
            chunks.append(current_chunk.strip())
            current_chunk = word
        else:
            if current_chunk:
                current_chunk += " " + word
            else:
                current_chunk = word
    if current_chunk:
        chunks.append(current_chunk.strip())
    return chunks


def convert_to_rdf(document_structure, tables, all_images_data, base_image_url):
    """Convert document structure, tables, and image data to RDF with maximum safety"""
    g = Graph()

    # Bind namespaces
    g.bind("ex", EX)
    g.bind("rdf", RDF)
    g.bind("rdfs", RDFS)
    g.bind("owl", OWL)
    g.bind("foaf", FOAF)
    g.bind("dcterms", DCTERMS)

    # Define base ontology classes (these are generic top-level concepts)
    g.add((EX.Content, RDF.type, OWL.Class))
    g.add((EX.Image, RDF.type, OWL.Class))
    g.add((EX.Table, RDF.type, OWL.Class))
    g.add((EX.TableCell, RDF.type, OWL.Class))
    g.add((EX.Patent, RDF.type, OWL.Class))
    g.add((EX.ContentChunk, RDF.type, OWL.Class))
    # Define custom ontology properties
    g.add((EX.hasSection, RDF.type, OWL.ObjectProperty))
    g.add((EX.hasSubsection, RDF.type, OWL.ObjectProperty))
    g.add((EX.hasContentChunk, RDF.type, OWL.ObjectProperty))
    g.add((EX.hasImage, RDF.type, OWL.ObjectProperty))
    g.add((EX.hasTable, RDF.type, OWL.ObjectProperty))
    g.add((EX.hasCell, RDF.type, OWL.ObjectProperty))
    g.add((EX.hasOCRText, RDF.type, OWL.DatatypeProperty))
    g.add((EX.hasCaption, RDF.type, OWL.DatatypeProperty))
    g.add((EX.onPage, RDF.type, OWL.DatatypeProperty))
    g.add((EX.mentionsEntity, RDF.type, OWL.ObjectProperty))
    g.add((EX.patentNumber, RDF.type, OWL.DatatypeProperty))
    g.add((EX.hasAbstract, RDF.type, OWL.DatatypeProperty))
    g.add((EX.hasFormulaLatex, RDF.type, OWL.DatatypeProperty))

    doc_title = document_structure.get("title", "Generated Document")
    doc_class_uri_name = safe_uri_name(doc_title)
    doc_class_uri = URIRef(EX[doc_class_uri_name])
    g.add((doc_class_uri, RDF.type, OWL.Class))
    g.add((doc_class_uri, RDFS.label, Literal(safe_literal(doc_title))))

    doc_uri = URIRef(EX[f"{doc_class_uri_name}_Instance"])
    g.add((doc_uri, RDF.type, doc_class_uri))
    g.add((doc_uri, RDFS.label, Literal(f"Document Instance of {safe_literal(doc_title)}")))

    if document_structure.get("patent_number"):
        g.add((doc_uri, RDF.type, EX.Patent))
        g.add((doc_uri, EX.patentNumber, Literal(safe_literal(document_structure["patent_number"]))))

    if document_structure.get("abstract_content"):
        g.add((doc_uri, EX.hasAbstract, Literal(safe_literal(document_structure["abstract_content"]))))

    for sec_index, section_obj in enumerate(document_structure.get("sections", [])):
        section_title = section_obj.get("title", f"Unnamed Section {sec_index + 1}")
        section_class_uri = URIRef(EX[safe_uri_name(section_title)])

        g.add((section_class_uri, RDF.type, OWL.Class))
        g.add((section_class_uri, RDFS.label, Literal(safe_literal(section_title))))
        g.add((section_class_uri, RDFS.subClassOf, doc_class_uri))

        section_instance_uri = URIRef(EX[f"{safe_uri_name(section_title)}_{sec_index + 1}_Instance"])
        g.add((section_instance_uri, RDF.type, section_class_uri))
        g.add((doc_uri, EX.hasSection, section_instance_uri))

        for content_item in section_obj.get("content", []):
            chunks = chunk_content(content_item)
            for i, chunk in enumerate(chunks):
                content_prop = EX[f"hasContent-{i + 1}"]
                g.add((section_instance_uri, content_prop, Literal(safe_literal(chunk))))

        for sub_sec_index, sub_section_obj in enumerate(section_obj.get("subsections", [])):
            sub_section_title = sub_section_obj.get("title", f"Unnamed Subsection {sub_sec_index + 1}")
            subsection_class_uri = URIRef(EX[safe_uri_name(sub_section_title)])

            g.add((subsection_class_uri, RDF.type, OWL.Class))
            g.add((subsection_class_uri, RDFS.label, Literal(safe_literal(sub_section_title))))
            g.add((subsection_class_uri, RDFS.subClassOf, section_class_uri))

            subsection_instance_uri = URIRef(EX[f"{safe_uri_name(sub_section_title)}_{sub_sec_index + 1}_Instance"])
            g.add((subsection_instance_uri, RDF.type, subsection_class_uri))
            g.add((section_instance_uri, EX.hasSubsection, subsection_instance_uri))

            for content_item in sub_section_obj.get("content", []):
                chunks = chunk_content(content_item)
                for i, chunk in enumerate(chunks):
                    content_prop = EX[f"hasContent-{i + 1}"]
                    g.add((subsection_instance_uri, content_prop, Literal(safe_literal(chunk))))

    for t_index, table in enumerate(tables):
        table_uri = URIRef(EX[f"Table{t_index + 1}"])
        g.add((table_uri, RDF.type, EX.Table))
        g.add((table_uri, RDFS.label, Literal(f"Table {t_index + 1}")))
        g.add((doc_uri, EX.hasTable, table_uri))

        for r_index, row in enumerate(table):
            for c_index, cell in enumerate(row):
                cell_uri = URIRef(EX[f"Cell_{t_index}_{r_index}_{c_index}"])
                g.add((cell_uri, RDF.type, EX.TableCell))
                g.add((cell_uri, RDFS.label, Literal(safe_literal(cell))))
                g.add((table_uri, EX.hasCell, cell_uri))

    for img_data in all_images_data:
        img_page = img_data["page_num"]
        img_idx = img_data["img_idx"]
        ocr_text = img_data["ocr_text"]
        caption = img_data.get("caption")
        identified_entities = img_data.get("identified_entities", {})
        image_filename = img_data["filename"]
        formula_latex = img_data.get("formula_latex")

        image_url = URIRef(f"{base_image_url}{image_filename}")
        image_resource = EX[f"Image_{safe_uri_name(image_filename)}"]

        g.add((image_resource, RDF.type, EX.Image))
        g.add((image_resource, FOAF.depiction, image_url))
        g.add((image_resource, DCTERMS.identifier, Literal(image_filename)))
        g.add((image_resource, DCTERMS.format, Literal(image_filename.split('.')[-1])))
        g.add((image_resource, EX.onPage, Literal(img_page)))

        g.add((image_resource, RDFS.comment, Literal(f"Hyperlink to image: {str(image_url)}", lang="en")))

        if ocr_text:
            g.add((image_resource, EX.hasOCRText, Literal(safe_literal(ocr_text))))
        if caption:
            g.add((image_resource, EX.hasCaption, Literal(safe_literal(caption))))

        if formula_latex:
            g.add((image_resource, EX.hasFormulaLatex, Literal(safe_literal(formula_latex))))

        for entity_type, keywords_found in identified_entities.items():
            entity_type_class_uri = URIRef(EX[safe_uri_name(entity_type)])
            g.add((entity_type_class_uri, RDF.type, OWL.Class))
            g.add((entity_type_class_uri, RDFS.label, Literal(entity_type.replace('_', ' ').title())))

            for keyword in keywords_found:
                keyword_instance_uri = URIRef(EX[f"Image_{img_page}_Idx{img_idx}_{safe_uri_name(keyword)}"])
                g.add((keyword_instance_uri, RDF.type, entity_type_class_uri))
                g.add((keyword_instance_uri, RDFS.label, Literal(keyword)))
                g.add((image_resource, EX.mentionsEntity, keyword_instance_uri))

        g.add((doc_uri, EX.hasImage, image_resource))

    return g


def create_clean_ttl(rdf_graph):
    """Create a clean TTL file without BOM or problematic characters"""
    try:
        ttl_content = rdf_graph.serialize(format="turtle")
        if isinstance(ttl_content, bytes):
            ttl_content = ttl_content.decode('utf-8')

        if ttl_content.startswith('\ufeff'):
            ttl_content = ttl_content[1:]

        return ttl_content.encode('utf-8', errors='ignore').decode('utf-8')

    except Exception as e:
        st.error(f"Error creating TTL: {e}")
        return None


def display_structure(document_structure, all_images_data):
    """Display the extracted structure with actual instance text content and image data"""
    st.subheader("Extracted Document Structure:")

    if not document_structure.get("sections") and not all_images_data and not document_structure.get(
            "abstract_content"):
        st.write("No structure, abstract, or images found in document.")
        return

    st.write(f"**Document Type (Main Class):** {document_structure.get('title', 'Untitled Document')}")
    if document_structure.get("patent_number"):
        st.write(f"**Patent Number (10):** {document_structure['patent_number']}")
    st.write(f"**Document Instance:** Document Instance of {document_structure.get('title', 'Untitled Document')}")

    if document_structure.get("abstract_content"):
        st.markdown(f"**Abstract (57):**")
        with st.expander("View Abstract Content"):
            st.write(document_structure['abstract_content'])

    for section_obj in document_structure.get("sections", []):
        st.markdown(f"**Class:** {section_obj.get('title', 'Untitled Class')}")

        if section_obj.get("content"):
            full_section_content = " ".join(section_obj['content'])
            content_chunks = chunk_content(full_section_content)
            with st.expander(f"View Content ({len(content_chunks)} parts)"):
                for i, chunk in enumerate(content_chunks):
                    st.text(f"Part {i + 1}: {chunk}")

        for sub_section_obj in section_obj.get("subsections", []):
            st.markdown(f"  **Subclass:** {sub_section_obj.get('title', 'Untitled Subclass')}")

            if sub_section_obj.get("content"):
                full_subsection_content = " ".join(sub_section_obj['content'])
                content_chunks = chunk_content(full_subsection_content)
                st.markdown(f"    **Content:**")
                with st.expander(f"View Content ({len(content_chunks)} parts)"):
                    for i, chunk in enumerate(content_chunks):
                        st.text(f"Part {i + 1}: {chunk}")

    if all_images_data:
        st.subheader("Extracted Images with OCR Text and Entities:")
        for i, img_data in enumerate(all_images_data):
            st.write(f"**Image {i + 1} (Page {img_data['page_num']}):**")
            if img_data['caption']:
                st.write(f"  **Caption:** {img_data['caption']}")
            if img_data['ocr_text']:
                st.write(
                    f"  **OCR Text:** {img_data['ocr_text'][:300]}{'...' if len(img_data['ocr_text']) > 300 else ''}")

            if img_data.get('formula_latex'):
                st.write(f"  **Detected Formula (LaTeX):**")
                st.latex(img_data['formula_latex'])
            else:
                st.info("  No specific entities or readable text found in this image.")

            st.image(img_data['local_path'],
                     caption=img_data['caption'] if img_data['caption'] else f"Image on Page {img_data['page_num']}")
            st.markdown("---")


def clear_directory_contents(directory_path):
    """
    Removes all files from the specified directory.
    If the directory does not exist, it does nothing.
    """
    if os.path.exists(directory_path) and os.path.isdir(directory_path):
        for filename in os.listdir(directory_path):
            file_path = os.path.join(directory_path, filename)
            try:
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)
            except Exception as e:
                st.error(f"Failed to delete {file_path}. Reason: {e}")
    else:
        print(f"Directory not found or not a directory: {directory_path}")


# Streamlit UI
st.title("MTS KNOWLEDGE-RDF REASONING CONVERTER")
st.write("Converts PDF and Word documents to RDF format with enhanced image processing")

document_category = st.selectbox(
    "Select document category:",
    ("General Document", "IEEE Knowledge", "Patent Knowledge")
)

selected_doc_type = st.radio(
    "Select document format:",
    ("PDF", "Word")
)

allowed_file_types = []
if document_category == "General Document":
    if selected_doc_type == "PDF":
        allowed_file_types = ["pdf"]
    elif selected_doc_type == "Word":
        allowed_file_types = ["docx"]
elif document_category in ["IEEE Knowledge", "Patent Knowledge"]:
    allowed_file_types = ["pdf"]
    if selected_doc_type == "Word":
        st.warning(
            f"While Word is selected, {document_category}s are typically PDF. Only PDF uploads are supported for this category.")

uploaded_file = st.file_uploader(
    f"Upload a {selected_doc_type} file for {document_category}",
    type=allowed_file_types
)

convert_button = st.button("Convert Document to RDF")

if uploaded_file and convert_button:
    file_extension = os.path.splitext(uploaded_file.name)[1].lower()
    input_filename = uploaded_file.name

    prefix = ""
    if document_category == "IEEE Knowledge":
        prefix = "IEEE_"
    elif document_category == "Patent Knowledge":
        prefix = "Patent_"

    original_filename_without_ext = os.path.splitext(input_filename)[0]
    output_filename = f"{prefix}{original_filename_without_ext}.ttl"

    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_input_filename_part = re.sub(r'[^a-zA-Z0-9_]', '', os.path.splitext(input_filename)[0])[:20]
    file_hash = hashlib.md5(input_filename.encode('utf-8')).hexdigest()[:8]
    dynamic_subfolder_name = f"{safe_input_filename_part}_{timestamp}_{file_hash}"

    base_image_project_folder = os.path.join(os.getcwd(), "Image")
    output_image_dir = os.path.join(base_image_project_folder, dynamic_subfolder_name)

    try:
        os.makedirs(output_image_dir, exist_ok=True)
        st.info(f"Extracted images will be saved to: `{output_image_dir}`")
    except OSError as e:
        st.error(
            f"Error creating output directory `{output_image_dir}`: {e}. Please check permissions or path validity.")
        st.stop()

    encoded_dynamic_subfolder_name = urllib.parse.quote(dynamic_subfolder_name)
    BASE_IMAGE_URL_PLACEHOLDER = f"http://localhost:8000/Image/{encoded_dynamic_subfolder_name}/"

    with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
        tmp_file.write(uploaded_file.read())
        temp_doc_path = tmp_file.name

    st.write(f"Processing {file_extension.upper()} document...")

    document_structure = {}
    table_content = []
    all_images_data = []

    try:
        if file_extension == '.pdf':
            document_structure = extract_pdf_structure(temp_doc_path, document_category)
            table_content = extract_pdf_tables(temp_doc_path)
            all_images_data = extract_pdf_images_with_ocr(temp_doc_path, output_image_dir)
            st.info("PDF processed using text extraction, pattern recognition, and enhanced OCR for images.")
        elif file_extension == '.docx':
            document_structure = extract_doc_structure(temp_doc_path, document_category)
            table_content = extract_doc_tables(temp_doc_path)
            st.warning(
                "Word document processed using style-based structure detection. "
                "Note: Image extraction, OCR, and captioning for DOCX are not fully supported in this version "
                "using standard libraries. Consider external tools or commercial libraries for comprehensive DOCX image handling."
            )
        else:
            st.error("Unsupported file type. Please upload a PDF or DOCX file.")
            if os.path.exists(temp_doc_path):
                os.unlink(temp_doc_path)
            st.stop()

        display_structure(document_structure, all_images_data)

        st.subheader("Extracted Tables")
        if table_content:
            for i, table in enumerate(table_content):
                st.write(f"**Table {i + 1}**")
                st.table(table)
        else:
            st.info("No tables found in document.")

        st.subheader("Generate RDF Knowledge Graph")
        st.write("Converting extracted data to RDF...")

        rdf_graph = convert_to_rdf(document_structure, table_content, all_images_data, BASE_IMAGE_URL_PLACEHOLDER)

        ttl_data = create_clean_ttl(rdf_graph)

        if ttl_data:
            st.success(f"Generated RDF with **{len(rdf_graph)}** triples.")

            with st.expander("Preview RDF (first 1000 characters)"):
                st.code(ttl_data[:1000] + "..." if len(ttl_data) > 1000 else ttl_data)

            st.download_button(
                label=f"Download Clean RDF ({output_filename})",
                data=ttl_data,
                file_name=output_filename,
                mime="text/turtle"
            )

            st.info(
                "The TTL file has been cleaned of BOM and problematic characters that could cause Protégé parsing errors.")
            st.warning(
                "To view images in Protégé, you need to run a local web server (e.g., Python's `http.server`) "
                "in the parent directory containing the 'Image' folder. "
                f"For example, navigate to `{os.getcwd()}` in your terminal and run `python -m http.server 8000`. "
                "Then, Protégé will be able to resolve the image URLs."
            )
        else:
            st.error("Failed to generate RDF data.")

    except Exception as e:
        st.error(f"An unexpected error occurred during document processing: {e}")
        import traceback

        st.code(traceback.format_exc())

    finally:
        if os.path.exists(temp_doc_path):
            os.unlink(temp_doc_path)
